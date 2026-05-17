"""Read and extract content from existing DOCX files."""

from pathlib import Path
from docx import Document


def read_docx(file_path: str) -> str:
    """Extract all text from a DOCX file."""
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def get_docx_info(file_path: str) -> dict:
    """Get metadata and statistics from a DOCX file."""
    doc = Document(file_path)
    core = doc.core_properties

    paragraphs = doc.paragraphs
    tables = doc.tables
    total_chars = sum(len(p.text) for p in paragraphs)

    return {
        "file": Path(file_path).name,
        "title": core.title or "",
        "author": core.author or "",
        "created": str(core.created) if core.created else "",
        "modified": str(core.modified) if core.modified else "",
        "paragraphs": len(paragraphs),
        "tables": len(tables),
        "characters": total_chars,
    }


def get_docx_paragraphs(file_path: str) -> list[dict]:
    """List all paragraphs with index, style, and text."""
    doc = Document(file_path)
    result = []
    for i, p in enumerate(doc.paragraphs):
        result.append({
            "index": i,
            "style": p.style.name if p.style else "Normal",
            "text": p.text,
        })
    return result


def get_docx_tables(file_path: str) -> list[dict]:
    """Extract tables with content as list of rows."""
    doc = Document(file_path)
    result = []
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        result.append({
            "index": i,
            "rows": len(table.rows),
            "columns": len(table.columns),
            "data": rows,
        })
    return result
